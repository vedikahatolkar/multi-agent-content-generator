from docx import Document

def create_word_file(content):

    doc = Document()

    doc.add_heading(
    'AI Generated Report',
    level=1
)

    doc.add_paragraph(content)

    file_path = "generated_report.docx"

    doc.save(file_path)

    return file_path